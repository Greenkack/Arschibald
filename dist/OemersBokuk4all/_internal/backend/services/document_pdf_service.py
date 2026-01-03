"""
Document PDF Conversion Service

This service provides functionality for converting various document formats to PDF bytes.
Supports Word documents, Excel spreadsheets, text documents, and multi-document merging.

Requirements: 14.8
Task: 228
"""

import io
import base64
from typing import List, Dict, Any, Optional, Union
from pathlib import Path
from datetime import datetime

try:
    from reportlab.lib.pagesizes import A4, letter
    from reportlab.lib.units import mm, cm
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, 
        PageBreak, Image as RLImage, KeepTogether
    )
    from reportlab.lib import colors
    from reportlab.pdfgen import canvas
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from PyPDF2 import PdfMerger, PdfReader, PdfWriter
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

try:
    import openpyxl
    from openpyxl.utils import get_column_letter
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False

try:
    from backend.core.pdf_bytes import PDFRenderingEngine, PDFMetadata
except ImportError:
    from core.pdf_bytes import PDFRenderingEngine, PDFMetadata


class DocumentConversionError(Exception):
    """Exception raised when document conversion fails"""
    pass


class DocumentPDFService:
    """
    Service for converting various document formats to PDF bytes.
    
    Supports:
    - Word documents (.docx)
    - Excel spreadsheets (.xlsx)
    - Text documents (.txt, .md)
    - Multi-document PDF merging
    """
    
    def __init__(self):
        self.pdf_engine = PDFRenderingEngine()
        self.styles = getSampleStyleSheet() if REPORTLAB_AVAILABLE else None
    
    def document_to_pdf_bytes(
        self,
        file_path: Optional[str] = None,
        file_content: Optional[bytes] = None,
        file_type: Optional[str] = None,
        metadata: Optional[PDFMetadata] = None
    ) -> bytes:
        """
        Convert any supported document to PDF bytes.
        
        Args:
            file_path: Path to the document file
            file_content: Raw bytes of the document (alternative to file_path)
            file_type: Document type ('docx', 'xlsx', 'txt', 'md')
            metadata: Optional PDF metadata
        
        Returns:
            bytes: PDF document as bytes
        
        Raises:
            DocumentConversionError: If conversion fails
            ValueError: If neither file_path nor file_content provided
        """
        if not REPORTLAB_AVAILABLE:
            raise ImportError("reportlab is required for PDF generation")
        
        # Determine file type
        if file_path:
            file_type = file_type or Path(file_path).suffix.lstrip('.')
        elif not file_type:
            raise ValueError("file_type must be provided when using file_content")
        
        # Route to appropriate converter
        if file_type in ['docx', 'doc']:
            return self.word_to_pdf_bytes(file_path, file_content, metadata)
        elif file_type in ['xlsx', 'xls']:
            return self.excel_to_pdf_bytes(file_path, file_content, metadata)
        elif file_type in ['txt', 'md', 'text']:
            return self.text_to_pdf_bytes(file_path, file_content, metadata)
        else:
            raise DocumentConversionError(f"Unsupported file type: {file_type}")
    
    def word_to_pdf_bytes(
        self,
        file_path: Optional[str] = None,
        file_content: Optional[bytes] = None,
        metadata: Optional[PDFMetadata] = None
    ) -> bytes:
        """
        Convert Word document to PDF bytes.
        
        Args:
            file_path: Path to .docx file
            file_content: Raw bytes of .docx file
            metadata: Optional PDF metadata
        
        Returns:
            bytes: PDF document as bytes
        """
        if not PYTHON_DOCX_AVAILABLE:
            raise ImportError("python-docx is required for Word document conversion")
        
        if not REPORTLAB_AVAILABLE:
            raise ImportError("reportlab is required for PDF generation")
        
        try:
            # Load Word document
            if file_path:
                doc = DocxDocument(file_path)
            elif file_content:
                doc = DocxDocument(io.BytesIO(file_content))
            else:
                raise ValueError("Either file_path or file_content must be provided")
            
            # Create PDF
            buffer = io.BytesIO()
            
            # Set default metadata if not provided
            if metadata is None:
                metadata = PDFMetadata(
                    title=Path(file_path).stem if file_path else "Word Document",
                    subject="Converted from Word document"
                )
            
            pdf_doc = self.pdf_engine.create_document(buffer, metadata)
            story = []
            
            # Process Word document content
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    # Determine style based on paragraph style
                    style = self._get_paragraph_style(paragraph)
                    story.append(Paragraph(paragraph.text, style))
                    story.append(Spacer(1, 6))
            
            # Process tables
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                
                if table_data:
                    pdf_table = self.pdf_engine.create_table(table_data)
                    story.append(pdf_table)
                    story.append(Spacer(1, 12))
            
            # Build PDF
            pdf_doc.build(story)
            
            buffer.seek(0)
            return buffer.getvalue()
            
        except Exception as e:
            raise DocumentConversionError(f"Failed to convert Word document: {str(e)}")
    
    def excel_to_pdf_bytes(
        self,
        file_path: Optional[str] = None,
        file_content: Optional[bytes] = None,
        metadata: Optional[PDFMetadata] = None,
        sheet_name: Optional[str] = None,
        max_rows: Optional[int] = None,
        max_cols: Optional[int] = None
    ) -> bytes:
        """
        Convert Excel spreadsheet to PDF bytes.
        
        Args:
            file_path: Path to .xlsx file
            file_content: Raw bytes of .xlsx file
            metadata: Optional PDF metadata
            sheet_name: Specific sheet to convert (default: active sheet)
            max_rows: Maximum number of rows to include
            max_cols: Maximum number of columns to include
        
        Returns:
            bytes: PDF document as bytes
        """
        if not OPENPYXL_AVAILABLE:
            raise ImportError("openpyxl is required for Excel document conversion")
        
        if not REPORTLAB_AVAILABLE:
            raise ImportError("reportlab is required for PDF generation")
        
        try:
            # Load Excel workbook
            if file_path:
                wb = openpyxl.load_workbook(file_path, data_only=True)
            elif file_content:
                wb = openpyxl.load_workbook(io.BytesIO(file_content), data_only=True)
            else:
                raise ValueError("Either file_path or file_content must be provided")
            
            # Select sheet
            if sheet_name:
                ws = wb[sheet_name]
            else:
                ws = wb.active
            
            # Create PDF
            buffer = io.BytesIO()
            
            # Set default metadata if not provided
            if metadata is None:
                metadata = PDFMetadata(
                    title=Path(file_path).stem if file_path else "Excel Spreadsheet",
                    subject=f"Converted from Excel - Sheet: {ws.title}"
                )
            
            pdf_doc = self.pdf_engine.create_document(buffer, metadata)
            story = []
            
            # Add title
            title_style = ParagraphStyle(
                'Title',
                parent=self.styles['Heading1'],
                fontSize=16,
                spaceAfter=20
            )
            story.append(Paragraph(f"Sheet: {ws.title}", title_style))
            story.append(Spacer(1, 12))
            
            # Extract data from worksheet
            table_data = []
            max_row = min(ws.max_row, max_rows) if max_rows else ws.max_row
            max_col = min(ws.max_column, max_cols) if max_cols else ws.max_column
            
            for row_idx in range(1, max_row + 1):
                row_data = []
                for col_idx in range(1, max_col + 1):
                    cell = ws.cell(row=row_idx, column=col_idx)
                    value = cell.value
                    
                    # Format numbers in German format
                    if isinstance(value, (int, float)):
                        value = self.pdf_engine.format_german_number(value)
                    elif value is None:
                        value = ""
                    else:
                        value = str(value)
                    
                    row_data.append(value)
                table_data.append(row_data)
            
            # Create table with styling
            if table_data:
                # Calculate column widths
                available_width = self.pdf_engine.width - 4*cm
                col_width = available_width / max_col
                col_widths = [col_width] * max_col
                
                pdf_table = self.pdf_engine.create_table(table_data, col_widths)
                story.append(pdf_table)
            
            # Build PDF
            pdf_doc.build(story)
            
            buffer.seek(0)
            return buffer.getvalue()
            
        except Exception as e:
            raise DocumentConversionError(f"Failed to convert Excel document: {str(e)}")
    
    def text_to_pdf_bytes(
        self,
        file_path: Optional[str] = None,
        file_content: Optional[bytes] = None,
        text_content: Optional[str] = None,
        metadata: Optional[PDFMetadata] = None,
        preserve_formatting: bool = True
    ) -> bytes:
        """
        Convert text document to PDF bytes.
        
        Args:
            file_path: Path to text file
            file_content: Raw bytes of text file
            text_content: Direct text content
            metadata: Optional PDF metadata
            preserve_formatting: Whether to preserve line breaks and spacing
        
        Returns:
            bytes: PDF document as bytes
        """
        if not REPORTLAB_AVAILABLE:
            raise ImportError("reportlab is required for PDF generation")
        
        try:
            # Get text content
            if text_content:
                text = text_content
            elif file_path:
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
            elif file_content:
                text = file_content.decode('utf-8')
            else:
                raise ValueError("One of text_content, file_path, or file_content must be provided")
            
            # Create PDF
            buffer = io.BytesIO()
            
            # Set default metadata if not provided
            if metadata is None:
                metadata = PDFMetadata(
                    title=Path(file_path).stem if file_path else "Text Document",
                    subject="Converted from text document"
                )
            
            pdf_doc = self.pdf_engine.create_document(buffer, metadata)
            story = []
            
            # Add title
            title_style = ParagraphStyle(
                'Title',
                parent=self.styles['Heading1'],
                fontSize=14,
                spaceAfter=20
            )
            story.append(Paragraph(metadata.title, title_style))
            story.append(Spacer(1, 12))
            
            # Process text content
            if preserve_formatting:
                # Preserve line breaks
                content_style = ParagraphStyle(
                    'PreformattedText',
                    parent=self.styles['Code'],
                    fontSize=10,
                    fontName='Courier',
                    leftIndent=0,
                    rightIndent=0
                )
                for line in text.split('\n'):
                    if line.strip():
                        story.append(Paragraph(line, content_style))
                    else:
                        story.append(Spacer(1, 6))
            else:
                # Flow text naturally
                content_style = ParagraphStyle(
                    'BodyText',
                    parent=self.styles['BodyText'],
                    fontSize=11,
                    alignment=TA_JUSTIFY
                )
                for paragraph in text.split('\n\n'):
                    if paragraph.strip():
                        story.append(Paragraph(paragraph, content_style))
                        story.append(Spacer(1, 12))
            
            # Build PDF
            pdf_doc.build(story)
            
            buffer.seek(0)
            return buffer.getvalue()
            
        except Exception as e:
            raise DocumentConversionError(f"Failed to convert text document: {str(e)}")
    
    def merge_pdf_documents(
        self,
        pdf_files: List[Union[str, bytes]],
        output_metadata: Optional[PDFMetadata] = None
    ) -> bytes:
        """
        Merge multiple PDF documents into a single PDF.
        
        Args:
            pdf_files: List of PDF file paths or PDF bytes
            output_metadata: Optional metadata for merged PDF
        
        Returns:
            bytes: Merged PDF document as bytes
        """
        if not PYPDF2_AVAILABLE:
            raise ImportError("PyPDF2 is required for PDF merging")
        
        try:
            merger = PdfMerger()
            
            # Add each PDF to merger
            for pdf_file in pdf_files:
                if isinstance(pdf_file, str):
                    # File path
                    merger.append(pdf_file)
                elif isinstance(pdf_file, bytes):
                    # PDF bytes
                    merger.append(io.BytesIO(pdf_file))
                else:
                    raise ValueError(f"Invalid PDF file type: {type(pdf_file)}")
            
            # Write merged PDF to buffer
            output_buffer = io.BytesIO()
            merger.write(output_buffer)
            merger.close()
            
            # Add metadata if provided
            if output_metadata:
                output_buffer.seek(0)
                reader = PdfReader(output_buffer)
                writer = PdfWriter()
                
                # Copy all pages
                for page in reader.pages:
                    writer.add_page(page)
                
                # Add metadata
                writer.add_metadata({
                    '/Title': output_metadata.title,
                    '/Author': output_metadata.author,
                    '/Subject': output_metadata.subject,
                    '/Creator': output_metadata.creator,
                    '/Keywords': ', '.join(output_metadata.keywords)
                })
                
                # Write with metadata
                final_buffer = io.BytesIO()
                writer.write(final_buffer)
                final_buffer.seek(0)
                return final_buffer.getvalue()
            
            output_buffer.seek(0)
            return output_buffer.getvalue()
            
        except Exception as e:
            raise DocumentConversionError(f"Failed to merge PDF documents: {str(e)}")
    
    def _get_paragraph_style(self, paragraph) -> ParagraphStyle:
        """
        Determine appropriate PDF style based on Word paragraph style.
        
        Args:
            paragraph: python-docx paragraph object
        
        Returns:
            ParagraphStyle: ReportLab paragraph style
        """
        if not self.styles:
            return None
        
        style_name = paragraph.style.name.lower()
        
        if 'heading 1' in style_name or 'title' in style_name:
            return self.styles['Heading1']
        elif 'heading 2' in style_name:
            return self.styles['Heading2']
        elif 'heading 3' in style_name:
            return self.styles['Heading3']
        else:
            return self.styles['BodyText']
    
    def convert_multiple_documents(
        self,
        documents: List[Dict[str, Any]],
        merge: bool = False,
        output_metadata: Optional[PDFMetadata] = None
    ) -> Union[List[bytes], bytes]:
        """
        Convert multiple documents to PDF.
        
        Args:
            documents: List of document dictionaries with keys:
                - file_path or file_content
                - file_type (optional)
                - metadata (optional)
            merge: Whether to merge all PDFs into one
            output_metadata: Metadata for merged PDF (if merge=True)
        
        Returns:
            List[bytes] if merge=False, bytes if merge=True
        """
        pdf_bytes_list = []
        
        for doc in documents:
            pdf_bytes = self.document_to_pdf_bytes(
                file_path=doc.get('file_path'),
                file_content=doc.get('file_content'),
                file_type=doc.get('file_type'),
                metadata=doc.get('metadata')
            )
            pdf_bytes_list.append(pdf_bytes)
        
        if merge:
            return self.merge_pdf_documents(pdf_bytes_list, output_metadata)
        else:
            return pdf_bytes_list


# Convenience functions

def word_to_pdf(
    file_path: str,
    output_path: Optional[str] = None,
    metadata: Optional[PDFMetadata] = None
) -> bytes:
    """
    Convert Word document to PDF.
    
    Args:
        file_path: Path to Word document
        output_path: Optional path to save PDF
        metadata: Optional PDF metadata
    
    Returns:
        bytes: PDF document as bytes
    """
    service = DocumentPDFService()
    pdf_bytes = service.word_to_pdf_bytes(file_path=file_path, metadata=metadata)
    
    if output_path:
        with open(output_path, 'wb') as f:
            f.write(pdf_bytes)
    
    return pdf_bytes


def excel_to_pdf(
    file_path: str,
    output_path: Optional[str] = None,
    metadata: Optional[PDFMetadata] = None,
    sheet_name: Optional[str] = None
) -> bytes:
    """
    Convert Excel spreadsheet to PDF.
    
    Args:
        file_path: Path to Excel file
        output_path: Optional path to save PDF
        metadata: Optional PDF metadata
        sheet_name: Specific sheet to convert
    
    Returns:
        bytes: PDF document as bytes
    """
    service = DocumentPDFService()
    pdf_bytes = service.excel_to_pdf_bytes(
        file_path=file_path,
        metadata=metadata,
        sheet_name=sheet_name
    )
    
    if output_path:
        with open(output_path, 'wb') as f:
            f.write(pdf_bytes)
    
    return pdf_bytes


def text_to_pdf(
    file_path: str,
    output_path: Optional[str] = None,
    metadata: Optional[PDFMetadata] = None
) -> bytes:
    """
    Convert text document to PDF.
    
    Args:
        file_path: Path to text file
        output_path: Optional path to save PDF
        metadata: Optional PDF metadata
    
    Returns:
        bytes: PDF document as bytes
    """
    service = DocumentPDFService()
    pdf_bytes = service.text_to_pdf_bytes(file_path=file_path, metadata=metadata)
    
    if output_path:
        with open(output_path, 'wb') as f:
            f.write(pdf_bytes)
    
    return pdf_bytes


def merge_pdfs(
    pdf_files: List[str],
    output_path: str,
    metadata: Optional[PDFMetadata] = None
) -> bytes:
    """
    Merge multiple PDF files into one.
    
    Args:
        pdf_files: List of PDF file paths
        output_path: Path to save merged PDF
        metadata: Optional metadata for merged PDF
    
    Returns:
        bytes: Merged PDF document as bytes
    """
    service = DocumentPDFService()
    merged_bytes = service.merge_pdf_documents(pdf_files, metadata)
    
    with open(output_path, 'wb') as f:
        f.write(merged_bytes)
    
    return merged_bytes
